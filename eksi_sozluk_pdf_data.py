'''
Ali Yurtseven

This program was written for downloding pdf pages from eksi sozluk by searching any variable entered by users!

target webpage= "www.eksisozluk.com"
'''

from selenium import webdriver
from bs4 import BeautifulSoup as soup
from urllib.request import urlopen as uReq
import os
from docx import Document
import urllib.error
import pathlib
import os
import pandas as pd


# Parameters should be entered carefully and parameters.txt file should be in the same directory with the code.

path_parameter_file=os.path.join(pathlib.Path(__file__).parent.absolute(),'parameters.txt')

if os.path.exists(path_parameter_file)==False:
    print('parameters.txt file should be in the same directory with the interactome_parameter.py file! There will be an error!')

data=pd.read_csv(path_parameter_file,delimiter='\t',header=None)

try:
    total_number_of_page=int(data[data.columns[0]][1])
    user_entered = data[data.columns[0]][0]
except Exception as p:
    print(p)

# Let me get the path where the user wants me to have
user=user_entered

# T
url=os.path.join('https://eksisozluk.com',user)

# Let me get the full pat, using Firefox browser
driver = webdriver.Firefox()
driver.get(url)
get_url = driver.current_url
driver.quit()

# Check how many pages do the page contains
counter=0
for i in range(1,100000000000000000000000):
    ne = get_url + "?p={}".format(i)
    try:
        conn = urllib.request.urlopen(ne)
    except urllib.error.HTTPError as e:
        # Return code error (e.g. 404, 501, ...)
        # ...
        nu = e.code
    except urllib.error.URLError as e:
        # Not an HTTP-specific error (e.g. connection refused)
        # ...
        nu = e.reason
    else:
        nu = 200
    if nu!=200:
        break
    counter+=1

page_number=counter
print('page number is: {}'.format(page_number))
# Lets find all texts, dates, and name of authors for all given pages!
texts=[]
date=[]
author=[]
if page_number>5:
    for i in range(page_number-(total_number_of_page-1),page_number+1):
        new_url=get_url+"?p={}".format(i)
        try:
            uClient = uReq(new_url)
            page = uClient.read()
            uClient.close()

            page_soup = soup(page, "html.parser")

            # Find the texts

            all_texts = page_soup.findAll("div",class_="content")
            for i in all_texts:
                texts.append(i.text)

            # Find the authors

            all_authors=page_soup.findAll("a",class_="entry-author")
            for j in all_authors:
                author.append(j.text)
            # Find the date

            all_dates = page_soup.findAll("a", class_="entry-date permalink")
            for k in all_dates:

                date.append(k.text)
        except:
            print("Your search is not found in the website please try again ")
elif 0<page_number<=5:
    for i in range(1, page_number + 1):
        new_url = get_url + "?p={}".format(i)
        try:
            uClient = uReq(new_url)
            page = uClient.read()
            uClient.close()

            page_soup = soup(page, "html.parser")

            # Find the texts

            all_texts = page_soup.findAll("div", class_="content")
            for i in all_texts:
                texts.append(i.text)

            # Find the authors

            all_authors = page_soup.findAll("a", class_="entry-author")
            for j in all_authors:
                author.append(j.text)
            # Find the date

            all_dates = page_soup.findAll("a", class_="entry-date permalink")
            for k in all_dates:
                date.append(k.text)
        except:
            print("Your search is not found in the website please try again ")

# Let me create the docx document!

def mydoc(texts,date,author,page_number):
    if page_number!=0:
        document = Document()

        document.add_heading(user.upper(), 0)

        p = document.add_paragraph(
            '"{}" icin "www.eksisozluk.com" adresinde yapilan yorumlar asagida verilmistir.'.format(user.upper()))

        for i in range(0, len(texts)):
            document.add_heading('Comment {}'.format(i + 1))
            document.add_paragraph('{}\t\t\t{}'.format(author[i], date[i]), style='Intense Quote')

            document.add_paragraph(texts[i][6:], style='Body Text')

        document.add_page_break()

        document.save('/home/a/Desktop/demo.docx')
    else:
        print("The website doesn't exist")
    return
mydoc(texts,date,author,page_number)

#